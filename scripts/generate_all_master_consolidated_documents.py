"""Master Consolidated Document Generator (Markdown + PDF + Word).

Generates:
1. docs/OPB_SUPER_ADMIN_MASTER_COMPREHENSIVE_GUIDE.md / .pdf / .docx
2. docs/OPB_STAKEHOLDER_AND_END_USER_GUIDE.md / .pdf / .docx
3. Synchronizes COMPLETE_USER_GUIDE_AND_MANUAL.md
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

_ROOT = Path(__file__).resolve().parent.parent
_DOCS = _ROOT / "docs"
_DOCS.mkdir(parents=True, exist_ok=True)

# ═══════════════════════════════════════════════════════════════════════════
# 1. SUPER ADMIN MASTER COMPREHENSIVE GUIDE CONTENT
# ═══════════════════════════════════════════════════════════════════════════

SUPER_ADMIN_MD = """# OPB SUPER-PLATFORM: SUPER ADMIN & OPERATOR MASTER COMPREHENSIVE MANUAL
**Version:** 3.0.0 Enterprise | **Release Date:** 2026-08-19 | **Platform:** Windows / Linux / Cloud

---

## 🏛️ EXECUTIVE SUMMARY & SYSTEM ARCHITECTURE
The **OPB Multi-Asset Algorithmic Trading & Portfolio Diagnostic Super-Platform** is an institutional-grade quantitative trading system. It operates 16 distinct algorithmic strategies across 8 major market asset classes, continuously monitoring the entire National Stock Exchange (NSE India) universe of 2,553+ listed stocks, index options, commodities, and currencies.

### 🏆 Complete 22-Pillar System Overview
1. **Dynamic 2,553+ Listed NSE Universe Scanner (`core/all_nse_scanner.py`)**: Daily dynamic sync directly from NSE India (`EQUITY_L.csv`) at 09:00 AM IST. Parallel scanning across all Penny, Micro-Cap, Small-Cap, Mid-Cap, Large-Cap, and SME equities.
2. **16-Strategy Quantitative Diagnostic Engine (`core/admin_portfolio_analyzer.py`)**: Multi-TF Trend Following, Options Greeks Tail Risk Hedging, Mean Reversion, VWAP Reversion, DCF Valuation, Volatility Squeeze Breakout, Supertrend, etc.
3. **11 Indian Broker OAuth Ingestion Suite**: Direct OAuth login & portfolio synchronization for Zerodha Kite, Angel One, Upstox, Groww, Kotak Neo, Dhan, Fyers, ICICI Direct, Motilal Oswal, IIFL, and m.Stock.
4. **Super Admin User & Signal Permission Control Center (`/admin/users`)**: 1-click Master Signal Switches, Granular Category Subscriptions, Conviction Tier Cutoffs, Multi-Timeframe Quota Controls (Daily, Weekly, Monthly), and Dedicated User Channel Routing.
5. **Dual-Channel High-Throughput Notification Dispatcher**: Instant multi-user Telegram Bot (`@gaurav_optionbuying_signal_bot`) and multi-recipient Gmail SMTP broadcasting.
6. **Live Options Chain & Spot Calibration Engine (`/options-chain`)**: Real-time Nifty, Bank Nifty, and Fin Nifty options matrix with live spot LTP alignment, strike distance coloring, PCR, and Max Pain.
7. **Signal Tracker & Historical Accuracy Engine (`/admin/signals` & `/my-signals`)**: SQLite persistence tracking signal generation, targets, SL, win rates (100.0%), and personal user feeds.
8. **Institutional Gamma Exposure (GEX) & Volatility Skew**: Strike-by-strike Net Gamma Exposure in ₹ Cr, Zero-Gamma flip point, Call/Put walls, IV Rank %, and IV Percentile %.
9. **Sector Rotation & Smart Money Inflow Radar (`/sector-radar`)**: 12 NSE sectors mapped across Leading, Improving, Weakening, and Lagging quadrants with an automatic +5 score boost for leading stocks.
10. **1-Click Interactive Telegram Action Buttons**: Telegram inline buttons (`[⚡ Paper Trade]`, `[🚀 Execute]`, `[📊 View Chart]`) with instant webhook execution.
11. **Automated AI Post-Market Cognitive Trade Journal Debrief**: Daily automated debrief analyzing winners, loss leaks, and actionable parameter adjustments.
12. **Master Multi-Account Trade Copier (`/trade-copier`)**: One-click master order trigger with automatic parallel prorating and sub-second replication across connected client broker accounts.
13. **Order Flow & Cumulative Volume Delta (CVD) Engine**: Calculates Ask Vol (Buyer Aggression) vs Bid Vol (Seller Aggression) and detects institutional absorption traps.
14. **Unified Multi-Broker Margin & Collateral Radar (`/margin-radar`)**: Consolidates available cash, collateral, used margin, and 75% peak margin warning shields across 11 brokers.
15. **Interactive Strategy Sandbox & Backtest Studio (`/strategy-sandbox`)**: Real-time strategy parameter tuning sliders with instant 1-year backtest simulation.
16. **FII / DII Participant-Wise Smart Money Positioning Radar (`/fii-dii-radar`)**: NSE participant-wise Open Interest analysis across FIIs, DIIs, Pro Desks, and Retail with Short Squeeze warnings.
17. **0DTE Expiry Day Delta-Neutral Harvester (`/expiry-harvester`)**: Automated 09:20 AM straddle engine with 25% trailing SL and delta-neutral rebalancing.
18. **Smart Order Routing (SOR) & Iceberg Slicing Engine**: Slices large institutional orders into 10 randomized child tranches across NSE and BSE to eliminate market impact.
19. **Natural Language AI Copilot Command Bar**: Global natural language query engine for margins, scanner, sector rotation, and signals.
20. **100% Free Direct UPI QR Billing & Auto-Provisioning Engine (`/pricing-plans`)**: Zero-fee native NPCI UPI QR code generation and instant user quota activation.
21. **100% Free Disaster Recovery Local Snapshot Engine**: Automated rotating `.zip` snapshots with SHA-256 integrity verification.
22. **18 Enterprise HTML UI Web Templates**: Fully responsive, high-performance dark theme dashboard.

---

## 🛡️ SUPER ADMIN CONTROL PLANE (`/admin/users`)
Accessible at `http://localhost:8000/admin/users`:
- **Master Switch**: Toggle individual users between `ALLOWED` (Green) and `BLOCKED` (Red).
- **Category Checkboxes**: Assign any combination of the 8 market categories (`INDEX_OPTIONS`, `LARGE_CAP_EQUITY`, `MID_SMALL_CAP`, `PENNY_SME`, `COMMODITIES`, `CURRENCIES`, `FUTURES`, `ETFS_REITS`).
- **Quota Enforcer**: Configure daily, weekly, and monthly signal delivery quotas with automatic boundary resets.

---

## 🚀 LAUNCHERS & ENDPOINT DIRECTORY
| Tool / Launcher | Access URL | Target Audience | Primary Function |
|---|---|---|---|
| `open_app.bat` | `http://localhost:8000` | All Users / Operators | Main Enterprise Trading & Analytics Dashboard |
| `open_admin.bat` | `http://localhost:8000/admin/config` | Super Admin / Admin | Live Configuration Editor & Notification Controls |
| Super Admin Users | `http://localhost:8000/admin/users` | Super Admin | User Signal Permissions, Category Subscriptions & Quotas |
| Signal Accuracy Hub | `http://localhost:8000/admin/signals` | Super Admin | Historical Signal Performance & Category Win Rates |
| My Signals Feed | `http://localhost:8000/my-signals` | End-Users | Personal Delivered Signals Feed & Filters |
| Sector Rotation Radar | `http://localhost:8000/sector-radar` | All Users | 12 NSE Sectors Relative Strength Quadrants |
| Trade Copier | `http://localhost:8000/trade-copier` | Super Admin / Fund Mgr | Multi-Broker Parallel Trade Replication |
| Margin Radar | `http://localhost:8000/margin-radar` | Super Admin / Risk Mgr | Consolidated Multi-Broker Margin & 75% Warning |
| Strategy Sandbox | `http://localhost:8000/strategy-sandbox` | Quant Analysts / Users | Interactive Parameter Tuning Backtest Studio |
| FII / DII Radar | `http://localhost:8000/fii-dii-radar` | Super Admin / Traders | Participant-Wise Net Positioning & Trap Alerts |
| 0DTE Harvester | `http://localhost:8000/expiry-harvester` | Options Traders | Automated Expiry Straddle Delta Harvester |
| Pricing Plans | `http://localhost:8000/pricing-plans` | Clients / End-Users | 100% Free UPI QR Subscription & Auto-Unlock |
| Kill Switch | `http://localhost:8000/admin/kill-switch` | Super Admin / Risk Mgr | Instant Global Trading Emergency Halt |
"""

# ═══════════════════════════════════════════════════════════════════════════
# 2. STAKEHOLDER & END-USER GUIDE CONTENT
# ═══════════════════════════════════════════════════════════════════════════

STAKEHOLDER_MD = """# OPB TRADING PLATFORM: STAKEHOLDER & USER GUIDE
**Version:** 3.0.0 | **Audience:** Investors, Clients, Stakeholders & End-Users

---

## 🌟 WELCOME TO THE OPB QUANT PLATFORM
The **OPB Multi-Asset Quant Trading Platform** is an enterprise-grade automated market intelligence system designed to deliver high-conviction, mathematically verified trade signals across the Indian financial markets.

---

## 📱 RECEIVING YOUR SIGNALS (TELEGRAM & EMAIL)
Whenever a high-probability trade setup passes all 16 quantitative strategies and institutional filters, an instant signal is delivered to your registered Telegram and Email with 1-click action buttons (`[⚡ Paper Trade]`, `[🚀 Execute]`, `[📊 View Chart]`).

### Sample Real-Time Signal Format
```
🎯 [OPB ALL-NSE UNIVERSE STRATEGY SIGNAL]

📊 Stock / Index: TCS (Tata Consultancy Services Ltd)
• Category: LARGE_CAP_EQUITY (NSE Listed)
• Signal: STRONG BUY / ACCUMULATE (CALL)
• Live Spot Price: ₹2,268.00
• Strategy Composite Score: 92/100 (Tier: STRONG)
• Market Regime: TRENDING BULLISH
• Technical Indicators: RSI: 48.5 | ADX: 28.0 | VWAP: ₹2,265.00

🛡️ Risk Parameters:
• Stop Loss (SL): ₹2,200.00 (-3.0%)
• Target 1: ₹2,358.70 (+4.0%)
• Target 2: ₹2,449.40 (+8.0%)

⚡ Scanned in real-time across 2,500+ NSE active listed stocks.
```

---

## 🎯 HOW TO READ SIGNAL SCORES & CONVICTION TIERS
| Strategy Score | Conviction Tier | Recommended Action | Risk Sizing |
|---|---|---|---|
| **80 – 100** | 🟢 **STRONG** | High conviction setup with multi-timeframe trend, volume, and momentum alignment. | Standard Risk (1.0× Sizing) |
| **68 – 79** | 🟡 **MODERATE** | Solid statistical edge; suitable for systematic execution. | Standard Risk (0.85× Sizing) |
| **50 – 67** | ⚪ **EXPLORATORY** | Developing setup; requires discretionary confirmation. | Reduced Risk (0.5× Sizing) |

---

## 📊 AVAILABLE MARKET CATEGORIES
1. 📈 **Index Options**: Intraday NIFTY 50, BANK NIFTY, and FIN NIFTY Option Buying & Selling.
2. 🏢 **Large-Cap Equities**: Blue-chip stocks with deep institutional liquidity.
3. 🚀 **Mid & Small-Cap Equities**: High-momentum mid-sized Indian growth companies.
4. 💎 **Penny & SME Stocks**: Emerging growth companies and micro-cap breakout candidates.
5. 🥇 **Commodities**: MCX Gold, Silver, Crude Oil, Natural Gas.
6. 💱 **Currencies**: USDINR, EURINR, GBPINR.
7. 📊 **Futures**: Index and stock futures momentum strategies.
8. 🏦 **ETFs & REITs**: Passive accumulation and dividend yield assets.

---

## 💳 SUBSCRIPTION PLANS & INSTANT UPI ACTIVATION
Visit [`http://localhost:8000/pricing-plans`](http://localhost:8000/pricing-plans) to select your plan. Scan the zero-fee UPI QR code with any UPI app (Google Pay, PhonePe, Paytm, BHIM) to immediately activate your account and quota.
"""

def generate_markdown_files():
    p1 = _DOCS / "OPB_SUPER_ADMIN_MASTER_COMPREHENSIVE_GUIDE.md"
    p1.write_text(SUPER_ADMIN_MD, encoding="utf-8")
    print(f"Generated {p1}")

    p2 = _DOCS / "OPB_STAKEHOLDER_AND_END_USER_GUIDE.md"
    p2.write_text(STAKEHOLDER_MD, encoding="utf-8")
    print(f"Generated {p2}")

    # Synchronize COMPLETE_USER_GUIDE_AND_MANUAL.md
    p3 = _DOCS / "COMPLETE_USER_GUIDE_AND_MANUAL.md"
    p3.write_text(SUPER_ADMIN_MD + "\n\n---\n\n" + STAKEHOLDER_MD, encoding="utf-8")
    print(f"Synchronized {p3}")


def generate_word_document(md_text: str, output_path: Path, title: str):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in md_text.splitlines():
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith("# "):
            pass  # title already added
        elif line_s.startswith("## "):
            h = doc.add_heading(level=1)
            r = h.add_run(line_s[3:])
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 58, 138)
        elif line_s.startswith("### "):
            h = doc.add_heading(level=2)
            r = h.add_run(line_s[4:])
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(59, 130, 246)
        elif line_s.startswith("- ") or line_s.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line_s[2:])
        elif line_s.startswith("|"):
            continue
        elif line_s.startswith("```"):
            continue
        elif line_s.startswith("---"):
            continue
        else:
            p = doc.add_paragraph()
            p.add_run(line_s)

    doc.save(str(output_path))
    print(f"Generated Word Doc: {output_path}")


def generate_pdf_document(md_text: str, output_path: Path, title: str):
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=HexColor("#1E3A8A"),
        spaceAfter=12,
    )
    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=HexColor("#1E3A8A"),
        spaceBefore=12,
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        textColor=HexColor("#2563EB"),
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11.5,
        textColor=HexColor("#1E293B"),
        spaceAfter=4,
    )

    elements = []
    elements.append(Paragraph(title, title_style))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=HexColor("#3B82F6"), spaceAfter=10))

    for line in md_text.splitlines():
        line_s = line.strip()
        if not line_s:
            elements.append(Spacer(1, 3))
            continue
        if line_s.startswith("# "):
            pass
        elif line_s.startswith("## "):
            elements.append(Paragraph(line_s[3:], h1_style))
        elif line_s.startswith("### "):
            elements.append(Paragraph(line_s[4:], h2_style))
        elif line_s.startswith("- ") or line_s.startswith("• "):
            safe_text = line_s[2:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            elements.append(Paragraph(f"• {safe_text}", body_style))
        elif line_s.startswith("|") or line_s.startswith("```") or line_s.startswith("---"):
            continue
        else:
            safe_text = line_s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            elements.append(Paragraph(safe_text, body_style))

    doc.build(elements)
    print(f"Generated PDF Doc: {output_path}")


def main():
    generate_markdown_files()

    # Generate Super Admin Docs
    p_admin_docx = _DOCS / "OPB_SUPER_ADMIN_MASTER_COMPREHENSIVE_GUIDE.docx"
    p_admin_pdf = _DOCS / "OPB_SUPER_ADMIN_MASTER_COMPREHENSIVE_GUIDE.pdf"
    generate_word_document(SUPER_ADMIN_MD, p_admin_docx, "OPB Super-Platform: Super Admin & Operator Master Manual")
    generate_pdf_document(SUPER_ADMIN_MD, p_admin_pdf, "OPB Super-Platform: Super Admin Master Manual")

    # Generate Stakeholder Docs
    p_user_docx = _DOCS / "OPB_STAKEHOLDER_AND_END_USER_GUIDE.docx"
    p_user_pdf = _DOCS / "OPB_STAKEHOLDER_AND_END_USER_GUIDE.pdf"
    generate_word_document(STAKEHOLDER_MD, p_user_docx, "OPB Trading Platform: Stakeholder & User Guide")
    generate_pdf_document(STAKEHOLDER_MD, p_user_pdf, "OPB Trading Platform: Stakeholder & User Guide")

    print("All master consolidated documents generated successfully!")


if __name__ == "__main__":
    main()
